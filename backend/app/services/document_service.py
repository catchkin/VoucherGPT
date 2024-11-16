from idlelib.iomenu import encoding
from typing import TypedDict, List, Optional, Dict, Any
import os
from fastapi import HTTPException, UploadFile, status
from sqlalchemy.ext.asyncio import AsyncSession
from datetime import datetime
import PyPDF2
from docx import Document as DocxDocument
import magic
import asyncio
from functools import partial
import json
from openai import AsyncOpenAI

from app.core.config import settings
from app.crud import crud_document, crud_section, crud_company
from app.models.document import DocumentType
from app.models.section import SectionType
from app.schemas.document import DocumentCreate, Document
from app.schemas.section import SectionCreate

import logging

# 로깅 설정
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)


class SectionData(TypedDict):
    type: SectionType
    title: str
    content: str
    order: int

class DocumentService:
    def __init__(self):
        """문서 처리 서비스 초기화"""
        self.supported_types = {
            'application/pdf': 'pdf',
            'application/msword': 'doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt'
        }
        self.client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)

    async def _extract_from_pdf(self, file_path: str) -> str:
        """
        PDF 파일에서 텍스트 추출
        """
        try:
            # PyPDF2는 동기식이므로 ThreadPool에서 실행
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, self._extract_from_pdf_sync, file_path)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error extracting text from PDF: {str(e)}"
            )

    def _extract_from_pdf_sync(self, file_path: str) -> str:
        """PDF 텍스트 추출을 위한 동기 메서드"""
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
        return '\n'.join(text_content)

    async def _extract_from_word(self, file_path: str) -> str:
        """Word 문서에서 텍스트 추출"""
        try:
            # python-docx도 동기식이므로 ThreadPool에서 실행
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, self._extract_from_word_sync, file_path)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error extracting text from Word document: {str(e)}"
            )

    def _extract_from_word_sync(self, file_path: str) -> str:
        """Word 텍스트 추출을 위한 동기 메서드"""
        doc = DocxDocument(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    async def _extract_from_txt(self, file_path: str) -> str:
        """텍스트 파일에서 텍스트 추출"""
        try:
            async with aiofiles.open(file_path, mode='r', encoding='utf-8') as file:
                return await file.read()
        except UnicodeDecodeError:
            # UTF-8로 읽기 실패시 다른 인코딩 시도
            try:
                async with aiofiles.open(file_path, mode='r', encoding='cp949') as file:
                    return await file.read()
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Error reading text file: {str(e)}"
                )

    async def extract_text_content(self, file_path: str) -> str:
        """파일 형식에 따라 적절한 텍스트 추출 메서드 호출"""
        # MIME 타입 감지
        mime_type = magic.from_file(file_path, mime=True)

        # 파일 확장자 추출
        file_ext = os.path.splitext(file_path)[1].lower()

        # MIME 타입과 확장자에 따른 추출 메서드 선택
        if mime_type == 'application/pdf' or file_ext == '.pdf':
            return await self._extract_from_pdf(file_path)
        elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'] \
                or file_ext in ['.doc', '.docx']:
            return await self._extract_from_word(file_path)
        elif mime_type == 'text/plain' or file_ext == '.txt':
            return await self._extract_from_txt(file_path)
        else:
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail=f"Unsupported file type: {mime_type}"
            )

    async def _check_file_corruption(self, file_path: str, mime_type: str) -> bool:
        """파일 무결성 검사"""
        try:
            if mime_type == 'application/pdf':
                with open(file_path, 'rb') as file:
                    PyPDF2.PdfReader(file)
            elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                DocxDocument(file_path)
            elif mime_type == 'text/plain':
                with open(file_path, 'r', encoding='utf-8') as file:
                    file.read()
            return True
        except Exception as e:
            return False

    async def save_uploaded_file(self, file: UploadFile) -> str:
        """
        업로드된 파일을 저장하고 저장 경로를 반환
        """
        if file.content_type not in self.supported_types:
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail=f"Unsupported file type. Supported types: {', '.join(self.supported_types.keys())}"
            )

        # 파일 저장 경로 생성
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        file_ext = self.supported_types[file.content_type]
        filename = f"{timestamp}_{file.filename}"
        file_path = os.path.join(settings.UPLOAD_DIR, filename)

        # 파일 저장
        try:
            os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
            with open(file_path, "wb+") as file_obj:
                content = await file.read()
                file_obj.write(content)
            return file_path
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error saving file: {str(e)}"
            )

    async def analyze_content(self, content: str) -> List[SectionData]:
        """문서 내용을 GPT를 사용하여 분석하고 섹션으로 분할"""

        prompt = f"""
        문서를 다음 섹션들로 분석하고 구조화해주세요:
        - Executive Summary (요약)
        - Company Overview (회사 개요)
        - Market Analysis (시장 분석)
        - Business Model (사업 모델)
        - Financial Plan (재무 계획)
        - Technical Description (기술 설명)

        각 섹션은 다음 형식의 JSON으로 반환해주세요:
        {{
            "sections": [
                {{
                    "type": "섹션 타입",
                    "title": "섹션 제목",
                    "content": "섹션 내용",
                    "order": 순서(0부터 시작)
                }}
            ]
        }}

        분석할 문서 내용:
        {content}
        """

        try:
            response = await self.client.chat.completions.create(
                model=settings.GPT_MODEL,
                messages=[
                    {"role": "system",
                     "content": "You are a document analyzer that specializes in business plans and company documents."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"}
            )

            return await self._process_gpt_response(response.choices[0].message.content)

        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error analyzing document content: {str(e)}"
            )

    async def _process_gpt_response(self, response_content: str) -> List[SectionData]:
        """GPT API 응답을 처리하여 섹션 데이터로 변환"""
        try:
            # JSON 파일
            if isinstance(response_content, str):
                data = json.loads(response_content)
            else:
                data = response_content

            sections = data.get('sections', [])
            processed_sections: List[SectionData] = []

            # 섹션 타입 매핑
            section_type_mapping = {
                'executive_summary': SectionType.EXECUTIVE_SUMMARY,
                'company_overview': SectionType.COMPANY_OVERVIEW,
                'market_analysis': SectionType.MARKET_ANALYSIS,
                'business_model': SectionType.BUSINESS_MODEL,
                'financial_plan': SectionType.FINANCIAL_PLAN,
                'technical_description': SectionType.TECHNICAL_DESCRIPTION
            }

            for idx, section in enumerate(sections):
                # 섹션 타입 변환
                section_type = section_type_mapping.get(
                    section['type'].lower().replace(' ', '_'),
                    SectionType.OTHER
                )

                processed_sections.append({
                    'type': section_type,
                    'title': section['title'],
                    'content': section['content'],
                    'order': section.get('order', idx)  # order가 없으면 인덱스 사용
                })

                return processed_sections

        except json.JSONDecodeError as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error parsing GPT response: {str(e)}"
            )
        except KeyError as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Missing required field in GPT response: {str(e)}"
            )

    async def generate_document_from_template(
        self,
        db: AsyncSession,
        template_id: int,
        company_id: int
    ) -> Document:
        """템플릿을 기반으로 새 문서 생성"""
        template = await crud_document.get_with_sections(db, id=template_id)
        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Template document not found"
            )

        # 회사 정보 조회
        company = await crud_company.get(db, id=company_id)
        if not company:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Company not found"
            )

        # 새 문서 생성
        new_document_data = DocumentCreate(
            title=f"{company.name} - {template.title}",
            type=template.type,
            content=template.content,
            company_id=company_id
        )

        new_document = await crud_document.create(db, obj_in=new_document_data)

        # 템플릿의 섹션들을 새 문서에 복사
        for template_section in template.sections:
            # GPT를 사용하여 섹션 내용을 회사에 맞게 수정
            customized_content = await self._customize_section_content(
                template_section.content,
                company
            )

            section_data = SectionCreate(
                type=template_section.type,
                title=template_section.title,
                content=customized_content,
                order=template_section.order,
                document_id=new_document.id,
                company_id=company_id
            )
            await crud_section.create(db, obj_in=section_data)

        return await crud_document.get_with_sections(db, id=new_document.id)

    async def _customize_section_content(
        self,
        template_content: str,
        company: Any
    ) -> str:
        """회사 정보를 바탕으로 섹션 내용 커스터마이징"""
        prompt = f"""
        다음 템플릿 내용을 주어진 회사 정보에 맞게 수정해주세요:

        템플릿 내용:
        {template_content}

        회사 정보:
        - 회사명: {company.name}
        - 업종: {company.industry}
        - 설립일: {company.establishment_date}
        - 직원 수: {company.employee_count}
        - 연간 매출: {company.annual_revenue}
        - 설명: {company.description}
        """

        try:
            response = await self.client.chat.completions.create(
                model=settings.GPT_MODEL,
                messages=[
                    {"role": "system",
                     "content": "You are an assistant that specializes in customizing business document content."},
                    {"role": "user", "content": prompt}
                ]
            )

            return response.choices[0].message.content.strip()

        except Exception as e:
            # GPT API 호출 실패 시 원본 내용 반환
            return template_content

    async def create_document_with_sections(
            self,
            db: AsyncSession,
            file: UploadFile,
            company_id: Optional[int] = None,
            document_type: DocumentType = DocumentType.BUSINESS_PLAN,
    ) -> Document:
        """
        문서 파일을 업로드하고 처리하여 섹션으로 분할하는 전체 프로세스 처리

        Args:
            db: 데이터베이스 세션
            file: 업로드된 파일
            company_id: 연관된 회사 ID (선택사항)
            document_type: 문서 유형 (기본값: 사업계획서)

        Returns:
            생성된 문서 (섹션 포함)
        """
        try:
            # 1. 파일 저장
            file_path = await self.save_uploaded_file(file)

            # 2. 파일 무결성 검사
            if not await self._check_file_corruption(file_path, file.content_type):
                os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Uploaded file is corrupted or invalid"
                )

            # 3. 텍스트 내용 추출
            try:
                content = await self.extract_text_content(file_path)
                if not content.strip():
                    raise ValueError("Extracted content is empty")
            except Exception as e:
                os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=f"Failed to extract content from file: {str(e)}"
                )

            # 4. 문서 메타데이터 생성
            try:
                document_in = DocumentCreate(
                    title=file.filename,
                    type=document_type,
                    content=content,
                    file_path=file_path,
                    file_name=file.filename,
                    mime_type=file.content_type,
                    company_id=company_id
                )
                document = await crud_document.create(db, obj_in=document_in)
            except Exception as e:
                os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to create document record: {str(e)}"
                )

            # 5. GPT를 사용한 내용 분석 및 섹션 생성
            try:
                sections_data = await self.analyze_content(content)

                for section_data in sections_data:
                    section_in = SectionCreate(
                        type=section_data["type"],
                        title=section_data["title"],
                        content=section_data["content"],
                        order=section_data["order"],
                        document_id=document.id,
                        company_id=company_id
                    )
                    await crud_section.create(db, obj_in=section_in)

            except Exception as e:
                # 실패 시 생성된 문서와 파일 정리
                await crud_document.remove(db, id=document.id)
                os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to analyze and create sections: {str(e)}"
                )

            # 6. 처리 결과 검증
            processed_document = await crud_document.get_with_sections(db, id=document.id)
            if not processed_document or not processed_document.sections:
                # 섹션이 없는 경우 문서와 파일 정리
                await crud_document.remove(db, id=document.id)
                os.remove(file_path)
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Failed to process document: No sections were created"
                )

            return processed_document

        except HTTPException:
            # 이미 처리된 HTTP 예외는 그대로 전달
            raise
        except Exception as e:
            if 'file_path' in locals() and os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Unexpected error during document processing: {str(e)}"
            )

    async def search_relevant_documents(
        self,
        db: AsyncSession,
        company_id: int,
        query: str
    ) -> List[Document]:
        """관련 문서 검색"""
        # 1. 회사 관련 문서
        company_docs = await crud_document.get_by_company(db, company_id=company_id)

        # 2. 학습용 문서 검색
        training_docs = await crud_document.get_by_type(
            db,
            doc_type=DocumentType.TRAINING_DATA
        )

        # 3. 연관성 점수 계산 및 필터링
        relevant_docs = self._filter_relevant_documents(
            query,
            company_docs + training_docs
        )

        return relevant_docs

    async def _filter_relevant_documents(
        self,
        query: str,
        documents: List[Document],
        threshold: float = 0.2, # 최소 연관성 점수
        max_documents: int = 5  # 최대 반환 문서 수
    ) -> List[Document]:
        """
        검색어와 문서들의 연관성을 평가하여 가장 관련성 높은 문서들을 반환

        Args:
             query: 검색어 또는 사용자 질문
             documents: 검색 대상 문서 리스트
             threshold: 최소 연관성 점수 (0~1)
             max_documents: 최대 반환 문서 수

        Returns:
            관련성 높은 문서 리스트
        """
        try:
            # 1. GPT 임베딩을 사용하여 검색어 벡터화
            query_embedding = await self._get_embedding(query)

            # 2. 각 문서의 연관성 점수 계산
            scored_documents = []
            for doc in documents:
                # 2.1 문서 내용에서 주요 부분 추출
                doc_content = self._extract_searchable_content(doc)

                # 2.2 문서 내용 임베딩
                doc_embedding = await self._get_embedding(doc_content)

                # 2.3 코사인 유사도 계산
                similarity = self._calculate_cosine_similarity(
                    query_embedding,
                    doc_embedding
                )

                # 2.4 추가 관련성 점수 계산 (문서 타입, 날짜 등 고려)
                relevance_score = self._calculate_relevance_score(
                    similarity,
                    doc
                )

                if relevance_score >= threshold:
                    scored_documents.append((doc, relevance_score))

            # 3. 점수순 정렬 및 상위 문서 선택
            scored_documents.sort(key=lambda x: x[1], reverse=True)
            return [doc for doc, score in scored_documents[:max_documents]]

        except Exception as e:
            logger.error(f"Error in filtering relevant documents: {str(e)}")
            # 오류 발생 시 기본 문서 반환
            return documents[:max_documents]

    async def _get_embedding(self, text: str) -> List[float]:
        """텍스트의 임베딩 벡터 생성"""
        try:
            response = await self.client.embeddings.create(
                model="text-embedding-ada-002",
                input=text
            )
            return response.data[0].embedding

        except Exception as e:
            logger.error(f"Error getting embedding: {str(e)}")
            raise

    def _extract_searchable_content(self, doc: Document) -> str:
        """문서에서 검색 가능한 주요 내용 추출"""
        # 문서 타입별 중요 섹션 추출
        content_parts = []

        # 제목 추가
        if doc.title:
            content_parts.append(doc.title)

        # 문서 타입별 처리
        if doc.type == DocumentType.BUSINESS_PLAN:
            # 사업계획서의 경우 주요 섹션 우선
            for section in doc.sections:
                if section.type in [
                    SectionType.EXECUTIVE_SUMMARY,
                    SectionType.MARKET_ANALYSIS,
                    SectionType.BUSINESS_MODEL
                ]:
                    content_parts.append(section.content)

        else:
            # 기타 문서는 전체 내용 사용
            if doc.content:
                content_parts.append(doc.content)

        # 최대 텍스트 길이 제한 (토큰 수 고려)
        combined_content = " ".join(content_parts)
        return combined_content[:8000]  # GPT 임베딩 모델 제한 고려

    def _calculate_cosine_similarity(
        self,
        vec1: List[float],
        vec2: List[float]
    ) -> float:
        """ 두 벡터 간의 코사인 유사도 계산"""
        dot_product = sum(a * b for a, b in zip(vec1, vec2))
        norm1 = sum(a * a for a in vec1) ** 0.5
        norm2 = sum(b * b for b in vec2) ** 0.5
        return dot_product / (norm1 * norm2) if norm1 * norm2 != 0 else 0

    def _calculate_relevance_score(
        self,
        similarity: float,
        doc: Document,
        base_weight: float = 0.7,
        type_weight: float = 0.2,
        date_weight: float = 0.1
    ) -> float:
        """
        최종 관련성 점수 계산
        - 코사인 유사도
        - 문서 타입 가중치
        - 최신성 가중치
        """
        # 1. 기본 유사도 점수
        score = similarity * base_weight

        # 2. 문서 타입 가중치
        type_score = self._get_document_type_score(doc.type)
        score += type_score * type_weight

        # 3. 최신성 가중치
        date_score = self._get_date_score(doc.created_at)
        score += date_score * date_weight

        return score

    def _get_document_type_score(self, doc_type: DocumentType) -> float:
        """문서 타입별 가중치 정의"""
        type_weights = {
            DocumentType.BUSINESS_PLAN: 1.0,  # 사업계획서 최우선
            DocumentType.COMPANY_PROFILE: 0.8,  # 회사 소개서 차순위
            DocumentType.TRAINING_DATA: 0.7,  # 학습 데이터
            DocumentType.FINANCIAL_REPORT: 0.6,  # 재무 보고서
            DocumentType.OTHER: 0.5  # 기타 문서
        }
        return type_weights.get(doc_type, 0.5)

    def _get_date_score(self, created_at: datetime) -> float:
        """문서 생성일자 기반 최신성 점수 계산"""
        now = datetime.utcnow()
        days_old = (now - created_at).days

        if days_old <= 30:  # 1개월 이내
            return 1.0
        elif days_old <= 90:  # 3개월 이내
            return 0.8
        elif days_old <= 180:  # 6개월 이내
            return 0.6
        elif days_old <= 365:  # 1년 이내
            return 0.4
        else:
            return 0.2


document_service = DocumentService()